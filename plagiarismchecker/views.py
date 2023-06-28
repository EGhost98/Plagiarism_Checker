from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algo import main
from docx import *
from plagiarismchecker.algo import fileSimilarity
import PyPDF2 

#home
def home(request):
    return render(request, 'pc/index.html') 

#web search(Text)
def test(request):
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())
    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text
    elif str(request.FILES['docfile']).endswith(".pdf"):
        pdfFileObj = open(request.FILES['docfile'], 'rb') 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 
        pageObj = pdfReader.getPage(0) 
        pdfFileObj.close() 
    percent,link = main.findSimilarity(value)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())
    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text
    result = fileSimilarity.findFileSimilarity(value1,value2)
    return render(request, 'pc/doc_compare.html',{'result': result})
